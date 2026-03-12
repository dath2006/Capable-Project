import os
import json
import tempfile
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, BackgroundTasks
from sqlalchemy.orm import Session
from typing import List

from ..database import get_db
from ..auth.dependencies import get_current_user
from ..auth.models import User
from .. import models, schemas
from ..services.rag import RAGService
from ..services.papers import PaperService
from ..ai import get_llm_config_error, has_llm_configuration

router = APIRouter(tags=["papers"])

@router.post("/papers/generate", response_model=schemas.QuestionPaperResponse)
async def generate_paper(
    payload: str = Form(...),
    file: UploadFile = File(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    if not has_llm_configuration():
        raise HTTPException(status_code=500, detail=get_llm_config_error())

    # Parse Payload
    try:
        config_data = json.loads(payload)
        config = schemas.PaperGenerateRequest(**config_data)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Invalid payload JSON: {str(e)}")

    if not file:
        raise HTTPException(status_code=400, detail="Document file is required")

    rag_service = RAGService()
    paper_service = PaperService()

    # Save uploaded file to temp file
    ext = os.path.splitext(file.filename)[1] if file.filename else ".txt"
    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as temp_file:
        content = await file.read()
        temp_file.write(content)
        temp_path = temp_file.name

    try:
        # Load and Index Document
        if file.filename.endswith(".pdf"):
            docs = rag_service.load_from_pdf(temp_path)
        elif file.filename.endswith(".docx"):
            docs = rag_service.load_from_docx(temp_path)
        elif file.filename.endswith(".txt") or file.filename.endswith(".md"):
            docs = rag_service.load_from_text(temp_path)
        else:
            raise HTTPException(status_code=400, detail="Unsupported file format")

        rag_service.index_documents(docs)

        # Retrieve relevant chunks based on topic list
        retrieved_docs = rag_service.similarity_search(" ".join(config.topic_focus) if config.topic_focus else "all concepts", k=30)
        if not retrieved_docs:
            raise HTTPException(status_code=400, detail="No extractable text found")
        chunks = [doc.page_content for doc in retrieved_docs]

    finally:
        if os.path.exists(temp_path):
            try:
                os.unlink(temp_path)
            except Exception:
                pass

    # 2. Save Paper DB Record
    db_paper = models.QuestionPaper(
        user_id=current_user.id,
        title=config.title,
        source_filename=file.filename,
        total_marks=config.total_marks,
        duration_minutes=config.duration_minutes,
        difficulty=config.difficulty,
        target_audience=config.target_audience
    )
    db.add(db_paper)
    db.commit()
    db.refresh(db_paper)

    # 3. Generate Sections
    import random
    order_idx = 0
    generated_marks = 0

    for sec_config in config.sections:
        db_section = models.PaperSection(
            paper_id=db_paper.id,
            section_type=sec_config.type,
            marks_per_question=sec_config.marks_per_question,
            order_index=order_idx
        )
        db.add(db_section)
        db.commit()
        db.refresh(db_section)
        order_idx += 1

        difficulties = paper_service._distribute_difficulty(config.difficulty, sec_config.count)

        for i in range(sec_config.count):
            chunk = random.choice(chunks)
            diff = difficulties[i]
            
            # Use appropriate service method
            try:
                if sec_config.type == "mcq":
                    res = paper_service.generate_mcq(chunk, diff)
                    db_q = models.Question(section_id=db_section.id, question_text=res.question, options=json.dumps(res.options), answer=res.options[res.correct_index], difficulty=diff, source_chunk=chunk, marks=sec_config.marks_per_question)
                    ans_text, ans_exp = res.options[res.correct_index], res.explanation
                elif sec_config.type == "true_false":
                    res = paper_service.generate_true_false(chunk, diff)
                    db_q = models.Question(section_id=db_section.id, question_text=res.statement, answer=res.answer, difficulty=diff, source_chunk=chunk, marks=sec_config.marks_per_question)
                    ans_text, ans_exp = res.answer, res.explanation
                elif sec_config.type == "fill_in_the_blank":
                    res = paper_service.generate_fill_in_the_blank(chunk, diff)
                    db_q = models.Question(section_id=db_section.id, question_text=res.sentence, answer=res.blank_word, difficulty=diff, source_chunk=chunk, marks=sec_config.marks_per_question)
                    ans_text, ans_exp = res.blank_word, res.context
                elif sec_config.type == "short_answer":
                    res = paper_service.generate_short_answer(chunk, diff)
                    db_q = models.Question(section_id=db_section.id, question_text=res.question, answer=res.ideal_answer, difficulty=diff, source_chunk=chunk, marks=sec_config.marks_per_question)
                    ans_text, ans_exp = res.ideal_answer, ", ".join(res.keywords)
                elif sec_config.type == "long_answer":
                    # For long answer, pick 3 random chunks to synthesize
                    synth_chunks = random.sample(chunks, min(3, len(chunks)))
                    res = paper_service.generate_long_answer(synth_chunks, diff)
                    db_q = models.Question(section_id=db_section.id, question_text=res.question, answer=res.ideal_answer, difficulty=diff, source_chunk="Multiple Chunks", marks=sec_config.marks_per_question)
                    ans_text, ans_exp = res.ideal_answer, res.marking_scheme
                elif sec_config.type == "case_study":
                    synth_chunks = random.sample(chunks, min(4, len(chunks)))
                    res = paper_service.generate_case_study(synth_chunks, diff)
                    passage_q = f"Passage:\n{res.passage}\n\nQuestions:\n" + "\n".join([f"{idx+1}. {q}" for idx, q in enumerate(res.sub_questions)])
                    db_q = models.Question(section_id=db_section.id, question_text=passage_q, answer=json.dumps(res.answers), difficulty=diff, source_chunk="Multiple Chunks", marks=sec_config.marks_per_question)
                    ans_text, ans_exp = "See JSON answers", "See JSON answers"
                else:
                    continue # Unknown type
            except Exception as ex:
                print(f"Generation error for {sec_config.type}: {ex}")
                continue
                
            db.add(db_q)
            db.commit()
            db.refresh(db_q)

            # Save Answer Key
            db_ans = models.AnswerKey(paper_id=db_paper.id, question_id=db_q.id, correct_answer=ans_text, explanation=ans_exp)
            db.add(db_ans)
            db.commit()
            
            generated_marks += sec_config.marks_per_question

    if generated_marks == 0:
        db.delete(db_paper)
        db.commit()
        raise HTTPException(status_code=500, detail="Failed to generate any questions")

    # Fetch complete tree structure to return
    db.refresh(db_paper)
    return db_paper

@router.get("/papers", response_model=List[schemas.QuestionPaperResponse])
def get_papers(db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    return db.query(models.QuestionPaper).filter(models.QuestionPaper.user_id == current_user.id).all()

@router.get("/papers/{paper_id}", response_model=schemas.QuestionPaperResponse)
def get_paper(paper_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    paper = db.query(models.QuestionPaper).filter(models.QuestionPaper.id == paper_id, models.QuestionPaper.user_id == current_user.id).first()
    if not paper:
        raise HTTPException(status_code=404, detail="Paper not found")
    return paper

@router.delete("/papers/{paper_id}")
def delete_paper(paper_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    paper = db.query(models.QuestionPaper).filter(models.QuestionPaper.id == paper_id, models.QuestionPaper.user_id == current_user.id).first()
    if not paper:
        raise HTTPException(status_code=404, detail="Paper not found")
    db.delete(paper)
    db.commit()
@router.post("/papers/{paper_id}/questions/{q_id}/regenerate", response_model=schemas.QuestionResponse)
def regenerate_question(paper_id: int, q_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    paper = db.query(models.QuestionPaper).filter(models.QuestionPaper.id == paper_id, models.QuestionPaper.user_id == current_user.id).first()
    if not paper:
        raise HTTPException(status_code=404, detail="Paper not found")

    db_q = db.query(models.Question).filter(models.Question.id == q_id, models.Question.section.has(paper_id=paper_id)).first()
    if not db_q:
        raise HTTPException(status_code=404, detail="Question not found")

    if not has_llm_configuration():
        raise HTTPException(status_code=500, detail=get_llm_config_error())

    paper_service = PaperService()

    sec_type = db_q.section.section_type
    chunk = db_q.source_chunk
    diff = db_q.difficulty

    try:
        if sec_type == "mcq":
            res = paper_service.generate_mcq(chunk, diff)
            db_q.question_text = res.question
            db_q.options = json.dumps(res.options)
            db_q.answer = res.options[res.correct_index]
            db_q.answer_key.correct_answer = db_q.answer
            db_q.answer_key.explanation = res.explanation
        elif sec_type == "true_false":
            res = paper_service.generate_true_false(chunk, diff)
            db_q.question_text = res.statement
            db_q.answer = res.answer
            db_q.answer_key.correct_answer = db_q.answer
            db_q.answer_key.explanation = res.explanation
        elif sec_type == "fill_in_the_blank":
            res = paper_service.generate_fill_in_the_blank(chunk, diff)
            db_q.question_text = res.sentence
            db_q.answer = res.blank_word
            db_q.answer_key.correct_answer = db_q.answer
            db_q.answer_key.explanation = res.context
        elif sec_type == "short_answer":
            res = paper_service.generate_short_answer(chunk, diff)
            db_q.question_text = res.question
            db_q.answer = res.ideal_answer
            db_q.answer_key.correct_answer = db_q.answer
            db_q.answer_key.explanation = ", ".join(res.keywords)
        elif sec_type == "long_answer":
            # For long answer we used multiple chunks previously
            res = paper_service.generate_long_answer([chunk], diff)
            db_q.question_text = res.question
            db_q.answer = res.ideal_answer
            db_q.answer_key.correct_answer = db_q.answer
            db_q.answer_key.explanation = res.marking_scheme
        elif sec_type == "case_study":
            res = paper_service.generate_case_study([chunk], diff)
            passage_q = f"Passage:\n{res.passage}\n\nQuestions:\n" + "\n".join([f"{idx+1}. {q}" for idx, q in enumerate(res.sub_questions)])
            db_q.question_text = passage_q
            db_q.answer = json.dumps(res.answers)
            db_q.answer_key.correct_answer = "See JSON answers"
            db_q.answer_key.explanation = "See JSON answers"
    except Exception as ex:
        raise HTTPException(status_code=500, detail=f"Regeneration failed: str({ex})")

    db.commit()
    db.refresh(db_q)
    return db_q

import io
from fastapi.responses import StreamingResponse

@router.post("/papers/{paper_id}/export")
def export_paper(paper_id: int, format: str = "pdf", type: str = "both", db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    paper = db.query(models.QuestionPaper).filter(models.QuestionPaper.id == paper_id, models.QuestionPaper.user_id == current_user.id).first()
    if not paper:
        raise HTTPException(status_code=404, detail="Paper not found")

    buffer = io.BytesIO()
    
    if format == "pdf":
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet
        
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        if type in ["paper", "both"]:
            story.append(Paragraph(paper.title, styles['Title']))
            story.append(Paragraph(f"Marks: {paper.total_marks} | Time: {paper.duration_minutes} min", styles['Normal']))
            story.append(Spacer(1, 12))
            
            for sec in paper.sections:
                story.append(Paragraph(f"Section - {sec.section_type.replace('_', ' ').title()}", styles['Heading2']))
                for q in sec.questions:
                    story.append(Paragraph(f"({q.marks} marks) {q.question_text}", styles['Normal']))
                    if sec.section_type == "mcq" and q.options:
                        opts = json.loads(q.options)
                        for idx, opt in enumerate(opts):
                            story.append(Paragraph(f"   {chr(65+idx)}. {opt}", styles['Normal']))
                    story.append(Spacer(1, 6))
            
        if type == "both":
            story.append(PageBreak())
            
        if type in ["answers", "both"]:
            story.append(Paragraph(f"{paper.title} - Answer Key", styles['Title']))
            story.append(Spacer(1, 12))
            
            for sec in paper.sections:
                story.append(Paragraph(f"Section - {sec.section_type.replace('_', ' ').title()}", styles['Heading2']))
                for q in sec.questions:
                    story.append(Paragraph(f"Q: {q.question_text}", styles['Normal']))
                    story.append(Paragraph(f"Ans: {q.answer_key.correct_answer}", styles['BodyText']))
                    if q.answer_key.explanation:
                        story.append(Paragraph(f"Exp: {q.answer_key.explanation}", styles['Italic']))
                    story.append(Spacer(1, 6))
                    
        doc.build(story)
        mimetype = "application/pdf"
        ext = "pdf"
        
    elif format == "docx":
        from docx import Document
        doc = Document()
        
        if type in ["paper", "both"]:
            doc.add_heading(paper.title, 0)
            doc.add_paragraph(f"Marks: {paper.total_marks} | Time: {paper.duration_minutes} min")
            
            for sec in paper.sections:
                doc.add_heading(f"Section - {sec.section_type.replace('_', ' ').title()}", level=1)
                for q in sec.questions:
                    doc.add_paragraph(f"({q.marks} marks) {q.question_text}")
                    if sec.section_type == "mcq" and q.options:
                        opts = json.loads(q.options)
                        for idx, opt in enumerate(opts):
                            doc.add_paragraph(f"   {chr(65+idx)}. {opt}")
                            
        if type == "both":
            doc.add_page_break()
            
        if type in ["answers", "both"]:
            doc.add_heading(f"{paper.title} - Answer Key", 0)
            for sec in paper.sections:
                doc.add_heading(f"Section - {sec.section_type.replace('_', ' ').title()}", level=1)
                for q in sec.questions:
                    doc.add_paragraph(f"Q: {q.question_text}")
                    doc.add_paragraph(f"Ans: {q.answer_key.correct_answer}", style='Intense Quote')
                    if q.answer_key.explanation:
                        doc.add_paragraph(f"Exp: {q.answer_key.explanation}")

        doc.save(buffer)
        mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ext = "docx"
    else:
        raise HTTPException(status_code=400, detail="Invalid format")

    buffer.seek(0)
    return StreamingResponse(buffer, media_type=mimetype, headers={"Content-Disposition": f"attachment; filename=paper_{paper_id}_{type}.{ext}"})
