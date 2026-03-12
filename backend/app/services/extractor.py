import io
import re
import pdfplumber
import docx


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using pdfplumber."""
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text.strip())
    raw = "\n\n".join(text_parts)
    return _clean_text(raw)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX file."""
    doc = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    raw = "\n\n".join(paragraphs)
    return _clean_text(raw)


def _clean_text(text: str) -> str:
    """Normalize whitespace and truncate to ~6000 words."""
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'\.{3,}', '...', text)
    text = text.strip()
    words = text.split()
    if len(words) > 6000:
        text = " ".join(words[:6000]) + "\n\n[Content truncated for processing]"
    return text